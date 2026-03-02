"""Word document builder for ECHR communicated cases."""

from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Colour palette (matches the reference doc) ─────────────────────────────
DARK_BLUE = RGBColor(0x1F, 0x39, 0x64)   # headings / title
MID_BLUE  = RGBColor(0x2E, 0x74, 0xB5)   # subtitle / articles line
BLACK     = RGBColor(0x00, 0x00, 0x00)


# ── Helpers ─────────────────────────────────────────────────────────────────

def _set_font(run, name="Aptos", size=12, bold=False, italic=False,
              color=BLACK):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def _add_paragraph(doc, text="", style=None, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph(text, style=style)
    p.alignment = alignment
    return p


def _add_page_break(doc):
    doc.add_page_break()


def _add_word_toc(doc):
    """Insert a Word TOC field (Heading 2–3, with hyperlinks and page numbers)."""
    p_head = doc.add_paragraph()
    run_h = p_head.add_run("TABLE OF CONTENTS")
    _set_font(run_h, size=16, bold=True, color=DARK_BLUE)
    p_head.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()

    def _wrap(child):
        r = OxmlElement("w:r")
        r.append(child)
        return r

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "2-3" \\z \\u '

    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")

    placeholder_r = OxmlElement("w:r")
    placeholder_t = OxmlElement("w:t")
    placeholder_t.text = (
        "Right-click here and choose \u201cUpdate Field\u201d to generate "
        "the table of contents."
    )
    placeholder_r.append(placeholder_t)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    p._p.append(_wrap(begin))
    p._p.append(_wrap(instr))
    p._p.append(_wrap(sep))
    p._p.append(placeholder_r)
    p._p.append(_wrap(end))


def _add_case_section(doc, case: dict):
    """Add a full case section (heading + body)."""
    # Case name as Heading 2 (picked up by the Word TOC field)
    p_name = doc.add_paragraph(style="Heading 2")
    p_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p_name.add_run(f"{case['docname'].upper()} ({case['appno']})")
    _set_font(run, size=14, bold=True, color=DARK_BLUE)
    p_name.paragraph_format.space_after = Pt(2)

    # Articles + keyword summary as Heading 3 (also picked up by the TOC)
    articles = case.get("article", "")
    summary = case.get("summary", "[Summary not available]")
    p_sub = doc.add_paragraph(style="Heading 3")
    sub_text = ""
    if articles:
        sub_text = f"Art. {articles}  –  "
    sub_text += summary
    run5 = p_sub.add_run(sub_text)
    _set_font(run5, size=11, italic=True, color=MID_BLUE)
    p_sub.paragraph_format.space_after = Pt(12)

    # Divider line (thin rule via border)
    _add_horizontal_rule(doc)

    # Subject matter body
    subject_matter = case.get("subject_matter", "").strip()
    if subject_matter:
        p_head = doc.add_paragraph()
        run_h = p_head.add_run("SUBJECT MATTER OF THE CASE")
        _set_font(run_h, size=11, bold=True, color=DARK_BLUE)

        for paragraph_text in subject_matter.split("\n"):
            stripped = paragraph_text.strip()
            if stripped:
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                _set_font(run, size=11)
                p.paragraph_format.space_after = Pt(4)

    # Questions to the parties
    questions = case.get("questions", "").strip()
    if questions:
        p_q_head = doc.add_paragraph()
        p_q_head.paragraph_format.space_before = Pt(10)
        run_qh = p_q_head.add_run("QUESTIONS TO THE PARTIES")
        _set_font(run_qh, size=11, bold=True, color=DARK_BLUE)

        for paragraph_text in questions.split("\n"):
            stripped = paragraph_text.strip()
            if stripped:
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                _set_font(run, size=11)
                p.paragraph_format.space_after = Pt(4)

    if case.get("has_appendix"):
        p_annex = doc.add_paragraph()
        p_annex.paragraph_format.space_before = Pt(10)
        run_annex = p_annex.add_run("Appendix omitted.")
        _set_font(run_annex, size=11, italic=True, color=BLACK)


def _add_horizontal_rule(doc):
    """Insert a thin horizontal rule paragraph."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E74B5")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(8)


# ── Public API ───────────────────────────────────────────────────────────────

def build_document(cases: list[dict], output_path: str) -> None:
    """
    Build a Word document summarising ECHR communicated cases.

    Args:
        cases: List of enriched case dicts (metadata + content + summary).
        output_path: File path for the output .docx file.
    """
    doc = Document()

    # ── Page margins ────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

    # ── Title page ──────────────────────────────────────────────────────────
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_title.add_run("SCAN OF COMMUNICATED CASES")
    _set_font(run_t, size=20, bold=True, color=DARK_BLUE)
    p_title.paragraph_format.space_before = Pt(72)
    p_title.paragraph_format.space_after  = Pt(6)

    # Subtitle: date range
    if cases:
        dates = [c.get("createddate", "") for c in cases if c.get("createddate")]
        if dates:
            try:
                first_dt = datetime.fromisoformat(min(dates).replace("Z", "+00:00"))
                last_dt  = datetime.fromisoformat(max(dates).replace("Z", "+00:00"))
                date_range_str = (
                    f"{first_dt.strftime('%d %B %Y').lstrip('0')} – "
                    f"{last_dt.strftime('%d %B %Y').lstrip('0')}"
                )
            except ValueError:
                date_range_str = ""
        else:
            date_range_str = ""
    else:
        date_range_str = ""

    if date_range_str:
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_sub = p_sub.add_run(date_range_str)
        _set_font(run_sub, size=14, italic=True, color=MID_BLUE)

    # Generated date
    p_gen = doc.add_paragraph()
    p_gen.alignment = WD_ALIGN_PARAGRAPH.CENTER
    today = datetime.today().strftime("%d %B %Y").lstrip("0")
    run_gen = p_gen.add_run(f"Generated: {today}")
    _set_font(run_gen, size=11, italic=True, color=BLACK)

    p_count = doc.add_paragraph()
    p_count.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_c = p_count.add_run(f"{len(cases)} case{'s' if len(cases) != 1 else ''}")
    _set_font(run_c, size=11, color=BLACK)

    _add_page_break(doc)

    # ── Table of Contents (Word field — updates automatically on open) ───────
    _add_word_toc(doc)

    _add_page_break(doc)

    # ── Case sections ───────────────────────────────────────────────────────
    for i, case in enumerate(cases):
        _add_case_section(doc, case)
        if i < len(cases) - 1:
            _add_page_break(doc)

    doc.save(output_path)
    print(f"Document saved: {output_path}")
